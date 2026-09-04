import os
import re
import uuid
import time
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple

BASE_DIR = Path(__file__).resolve().parent.parent
UPLOAD_DIR = BASE_DIR / "uploads"
GENERATED_DIR = BASE_DIR / "generated_docs"

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
GENERATED_DIR.mkdir(parents=True, exist_ok=True)

class DocumentService:
    """
    Handles local Document Ingestion (PDF, Word, Code, Text, CSV)
    and High-Fidelity Document Generation (PDF and Word .docx).
    """

    def __init__(self):
        self.upload_dir = UPLOAD_DIR
        self.generated_dir = GENERATED_DIR

    # -------------------------------------------------------------
    # 1. Document Text Ingestion
    # -------------------------------------------------------------
    def extract_text(self, file_path: Path, filename: str) -> Dict[str, Any]:
        """
        Extracts clean text content from various file formats (.pdf, .docx, .txt, .csv, .md, .py, .json).
        """
        ext = file_path.suffix.lower()
        extracted_text = ""
        meta: Dict[str, Any] = {"filename": filename, "extension": ext}

        if ext == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(str(file_path))
                meta["page_count"] = len(reader.pages)
                pages_text = []
                for i, page in enumerate(reader.pages):
                    pt = page.extract_text() or ""
                    if pt.strip():
                        pages_text.append(f"[Page {i+1}]\n{pt.strip()}")
                extracted_text = "\n\n".join(pages_text)
            except Exception as e:
                extracted_text = f"Error reading PDF: {str(e)}"

        elif ext in [".docx", ".doc"]:
            try:
                import docx
                doc = docx.Document(str(file_path))
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells if c.text.strip()]
                        if cells:
                            tables_text.append(" | ".join(cells))
                all_parts = paragraphs + tables_text
                meta["paragraph_count"] = len(paragraphs)
                extracted_text = "\n\n".join(all_parts)
            except Exception as e:
                extracted_text = f"Error reading Word document: {str(e)}"

        else:
            # Plain text, Markdown, CSV, Python, JSON, etc.
            try:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    extracted_text = f.read()
            except Exception as e:
                extracted_text = f"Error reading text file: {str(e)}"

        # Context truncation safeguard: limit to ~40,000 characters to prevent overflow
        truncated = False
        if len(extracted_text) > 40000:
            extracted_text = extracted_text[:40000] + "\n\n...[Document truncated for context limits]..."
            truncated = True

        meta["char_count"] = len(extracted_text)
        meta["word_count"] = len(extracted_text.split())
        meta["truncated"] = truncated

        return {
            "text": extracted_text,
            "metadata": meta
        }

    # -------------------------------------------------------------
    # 2. PDF Document Generation (ReportLab)
    # -------------------------------------------------------------
    def generate_pdf(self, title: str, markdown_content: str, filename: Optional[str] = None) -> Dict[str, Any]:
        """
        Builds a styled PDF document using ReportLab.
        Applies Claude-inspired styling: terracotta accents, clean typography, and margins.
        """
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

        clean_slug = re.sub(r'[^a-zA-Z0-9_\-]', '_', title[:30].strip().lower())
        out_name = filename or f"{clean_slug}_{int(time.time())}.pdf"
        if not out_name.endswith(".pdf"):
            out_name += ".pdf"
        out_path = self.generated_dir / out_name

        doc = SimpleDocTemplate(
            str(out_path),
            pagesize=letter,
            leftMargin=54,
            rightMargin=54,
            topMargin=54,
            bottomMargin=54
        )

        styles = getSampleStyleSheet()

        # Custom Claude-inspired styles
        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=22,
            leading=26,
            textColor=colors.HexColor('#da7756'),
            spaceAfter=12
        )

        h1_style = ParagraphStyle(
            'Heading1_Custom',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=15,
            leading=19,
            textColor=colors.HexColor('#22211e'),
            spaceBefore=14,
            spaceAfter=8
        )

        h2_style = ParagraphStyle(
            'Heading2_Custom',
            parent=styles['Heading3'],
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=16,
            textColor=colors.HexColor('#4a4640'),
            spaceBefore=10,
            spaceAfter=6
        )

        body_style = ParagraphStyle(
            'Body_Custom',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=15,
            textColor=colors.HexColor('#2a2825'),
            spaceAfter=8
        )

        bullet_style = ParagraphStyle(
            'Bullet_Custom',
            parent=body_style,
            leftIndent=15,
            firstLineIndent=-10,
            spaceAfter=4
        )

        meta_style = ParagraphStyle(
            'Meta_Custom',
            parent=styles['Italic'],
            fontName='Helvetica-Oblique',
            fontSize=9,
            leading=12,
            textColor=colors.HexColor('#888278'),
            spaceAfter=14
        )

        story = []

        # Document Header
        story.append(Paragraph(self._escape_xml(title), title_style))
        story.append(Paragraph(f"Generated by Nexus AI &bull; {time.strftime('%B %d, %Y')}", meta_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e5e0d8'), spaceAfter=14))

        # Parse Markdown content into ReportLab Flowables
        lines = markdown_content.split("\n")
        in_code_block = False
        code_lines = []

        for line in lines:
            stripped = line.strip()

            if stripped.startswith("```"):
                if in_code_block:
                    in_code_block = False
                    code_text = "<br/>".join([self._escape_xml(c) for c in code_lines])
                    story.append(Paragraph(f"<font face='Courier' size=8>{code_text}</font>", body_style))
                    story.append(Spacer(1, 6))
                    code_lines = []
                else:
                    in_code_block = True
                    code_lines = []
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            if not stripped:
                story.append(Spacer(1, 6))
                continue

            if stripped.startswith("# "):
                text = stripped[2:].strip()
                story.append(Paragraph(self._format_inline_markdown(text), h1_style))
            elif stripped.startswith("## "):
                text = stripped[3:].strip()
                story.append(Paragraph(self._format_inline_markdown(text), h2_style))
            elif stripped.startswith("### "):
                text = stripped[4:].strip()
                story.append(Paragraph(self._format_inline_markdown(text), h2_style))
            elif stripped.startswith("- ") or stripped.startswith("* "):
                text = stripped[2:].strip()
                story.append(Paragraph(f"&bull; {self._format_inline_markdown(text)}", bullet_style))
            elif re.match(r'^\d+\.\s', stripped):
                text = re.sub(r'^\d+\.\s', '', stripped).strip()
                story.append(Paragraph(f"&bull; {self._format_inline_markdown(text)}", bullet_style))
            else:
                story.append(Paragraph(self._format_inline_markdown(stripped), body_style))

        doc.build(story)

        file_size = out_path.stat().st_size
        return {
            "filename": out_name,
            "format": "pdf",
            "file_path": str(out_path),
            "download_url": f"/api/download/{out_name}",
            "size_bytes": file_size,
            "title": title
        }

    # -------------------------------------------------------------
    # 3. Word Document (.docx) Generation
    # -------------------------------------------------------------
    def generate_docx(self, title: str, markdown_content: str, filename: Optional[str] = None) -> Dict[str, Any]:
        """
        Builds a styled Word (.docx) document using python-docx.
        """
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        clean_slug = re.sub(r'[^a-zA-Z0-9_\-]', '_', title[:30].strip().lower())
        out_name = filename or f"{clean_slug}_{int(time.time())}.docx"
        if not out_name.endswith(".docx"):
            out_name += ".docx"
        out_path = self.generated_dir / out_name

        doc = docx.Document()

        # Configure 1-inch margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(4)
        run_title = p_title.add_run(title)
        run_title.font.name = 'Calibri'
        run_title.font.size = Pt(22)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(218, 119, 86)  # Claude terracotta

        # Subtitle / Timestamp
        p_meta = doc.add_paragraph()
        p_meta.paragraph_format.space_after = Pt(16)
        run_meta = p_meta.add_run(f"Generated by Nexus AI • {time.strftime('%B %d, %Y')}")
        run_meta.font.name = 'Calibri'
        run_meta.font.size = Pt(9.5)
        run_meta.font.italic = True
        run_meta.font.color.rgb = RGBColor(140, 134, 126)

        # Parse Markdown lines
        lines = markdown_content.split("\n")
        in_code_block = False
        code_lines = []

        for line in lines:
            stripped = line.strip()

            if stripped.startswith("```"):
                if in_code_block:
                    in_code_block = False
                    p_code = doc.add_paragraph()
                    p_code.paragraph_format.left_indent = Inches(0.25)
                    p_code.paragraph_format.space_after = Pt(8)
                    run_code = p_code.add_run("\n".join(code_lines))
                    run_code.font.name = 'Courier New'
                    run_code.font.size = Pt(9)
                    run_code.font.color.rgb = RGBColor(70, 68, 64)
                    code_lines = []
                else:
                    in_code_block = True
                    code_lines = []
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            if not stripped:
                continue

            if stripped.startswith("# "):
                h = doc.add_heading(level=1)
                h.paragraph_format.space_before = Pt(14)
                h.paragraph_format.space_after = Pt(4)
                run = h.add_run(stripped[2:].strip())
                run.font.name = 'Calibri'
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(34, 33, 30)

            elif stripped.startswith("## "):
                h = doc.add_heading(level=2)
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(4)
                run = h.add_run(stripped[3:].strip())
                run.font.name = 'Calibri'
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(74, 70, 64)

            elif stripped.startswith("### "):
                h = doc.add_heading(level=3)
                h.paragraph_format.space_before = Pt(10)
                h.paragraph_format.space_after = Pt(2)
                run = h.add_run(stripped[4:].strip())
                run.font.name = 'Calibri'
                run.font.size = Pt(11.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor(100, 95, 88)

            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(3)
                self._add_formatted_runs_docx(p, stripped[2:].strip())

            elif re.match(r'^\d+\.\s', stripped):
                clean_text = re.sub(r'^\d+\.\s', '', stripped).strip()
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.space_after = Pt(3)
                self._add_formatted_runs_docx(p, clean_text)

            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                self._add_formatted_runs_docx(p, stripped)

        doc.save(str(out_path))

        file_size = out_path.stat().st_size
        return {
            "filename": out_name,
            "format": "docx",
            "file_path": str(out_path),
            "download_url": f"/api/download/{out_name}",
            "size_bytes": file_size,
            "title": title
        }

    # -------------------------------------------------------------
    # 4. Intent Detection: PDF / Word Document Requests
    # -------------------------------------------------------------
    def detect_doc_request(self, prompt: str) -> Optional[Tuple[str, str]]:
        """
        Detects if the user prompt is asking to produce/generate a downloadable Word or PDF file.
        Returns ('pdf', topic) or ('docx', topic), or None.
        """
        text = prompt.strip().lower()

        # Check for Word / DOCX requests
        word_patterns = [
            r'(?:generate|create|make|export|download|write)\s+(?:me\s+)?(?:a\s+)?(?:word\s+doc(?:ument)?|docx|word\s+file)\s*(?:about|on|for|of)?\s*(.*)',
            r'(?:export|save|download)\s+(?:this\s+)?(?:as\s+)?(?:a\s+)?(?:word\s+doc(?:ument)?|docx|word\s+file)',
        ]
        for p in word_patterns:
            m = re.search(p, text, re.IGNORECASE)
            if m:
                topic = m.group(1).strip() if m.groups() and m.group(1) else "Document"
                return ("docx", topic or "Word Document")

        # Check for PDF requests
        pdf_patterns = [
            r'(?:generate|create|make|export|download|write)\s+(?:me\s+)?(?:a\s+)?(?:pdf\s+doc(?:ument)?|pdf\s+file|pdf)\s*(?:about|on|for|of)?\s*(.*)',
            r'(?:export|save|download)\s+(?:this\s+)?(?:as\s+)?(?:a\s+)?(?:pdf\s+doc(?:ument)?|pdf\s+file|pdf)',
        ]
        for p in pdf_patterns:
            m = re.search(p, text, re.IGNORECASE)
            if m:
                topic = m.group(1).strip() if m.groups() and m.group(1) else "Document"
                return ("pdf", topic or "PDF Document")

        return None

    # -------------------------------------------------------------
    # Helper formatters
    # -------------------------------------------------------------
    @staticmethod
    def _escape_xml(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    @classmethod
    def _format_inline_markdown(cls, text: str) -> str:
        escaped = cls._escape_xml(text)
        # Bold: **text** -> <b>text</b>
        escaped = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', escaped)
        # Italic: *text* -> <i>\1</i>
        escaped = re.sub(r'(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)', r'<i>\1</i>', escaped)
        # Code: `code` -> <font face="Courier">\1</font>
        escaped = re.sub(r'`(.*?)`', r"<font face='Courier' color='#da7756'>\1</font>", escaped)
        return escaped

    @staticmethod
    def _add_formatted_runs_docx(paragraph, text: str):
        """
        Parses bold **text** and italic *text* into docx runs.
        """
        pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
        tokens = pattern.split(text)

        for token in tokens:
            if not token:
                continue
            if token.startswith("**") and token.endswith("**"):
                r = paragraph.add_run(token[2:-2])
                r.bold = True
            elif token.startswith("*") and token.endswith("*"):
                r = paragraph.add_run(token[1:-1])
                r.italic = True
            elif token.startswith("`") and token.endswith("`"):
                r = paragraph.add_run(token[1:-1])
                r.font.name = 'Courier New'
                r.font.size = Pt(9.5)
            else:
                paragraph.add_run(token)

document_service = DocumentService()
